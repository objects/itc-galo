"""Deteccion y extraccion de actos normativos + registro del corpus consolidado (F4).

Los actos (decretos/resoluciones) que reglamentan o modifican el Decreto 555/2021
se integran al corpus consolidado como documentos nuevos versionados en git
(FR-013). Esta capa es la frontera de ingesta de F4 (data-model.md:15-18):
detecta el formato del archivo fuente (FR-001), extrae los articulos (sisjur
reutilizado D4 o extraccion generica D5), valida la relacion con el 555
(FR-014) y escribe el JSONL + `.sha256` + registro con deduplicacion por hash
(FR-007) y fallo atomico por documento (FR-009, SC-006).

Los errores son tipificados propios del CLI (`ErrorIngesta`): la taxonomia de
10 codigos de `app/errores.py` NO se modifica (FR-011).
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Any

from app.ingesta.corpus import (
    _extraer_banner_derogacion,
    _extraer_derogados,
    _extraer_upls,
    hash_documento,
    parsear_articulos,
)
from app.models import (
    ArticuloNormativo,
    DocumentoNormativo,
    FECHA_VIGENCIA_555,
    FormatoDocumento,
)

# Formato del archivo fuente por extension (FR-001); los magic bytes mandan
# sobre la extension (D5: el contenido es la senal autoritativa).
FORMATOS_POR_EXTENSION: dict[str, FormatoDocumento] = {
    ".html": "sisjur_html",
    ".htm": "sisjur_html",
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "txt",
}

# Enlace interno sisjur a articulos del Decreto 555 (H2): `Norma1.jsp?i=119582#NNN`.
# Es la referencia verificable por maquina para `relacion_con_555` (FR-014).
ID_SISJUR_555 = "119582"
ENLACE_ARTICULO_555_PATRON = re.compile(rf"Norma1\.jsp\?i={ID_SISJUR_555}#(\d+)")

# Patron generico de articulo para PDF/DOCX/Markdown/TXT (D5): "Artículo N."
# con numero arabigo u ordinal textual (Primero -> 1, Unico -> 1). El ancla al
# inicio de linea evita falsos positivos ("...articulo 12 de la Ley X").
ARTICULO_GENERICO_PATRON = re.compile(
    r"^\s*Art[ií]culo\s+(?:(?P<numero>\d{1,3})\s*[º°]?\.?\s*[º°]?|(?P<ordinal>[^\W\d_]+))\.?\s*",
    re.IGNORECASE | re.MULTILINE,
)

ORDINALES_A_NUMERO: dict[str, int] = {
    "único": 1,
    "unico": 1,
    "primero": 1,
    "primera": 1,
    "prime": 1,
    "segundo": 2,
    "segunda": 2,
    "tercero": 3,
    "tercera": 3,
    "tercer": 3,
    "cuarto": 4,
    "cuarta": 4,
    "quinto": 5,
    "quinta": 5,
    "sexto": 6,
    "sexta": 6,
    "séptimo": 7,
    "séptima": 7,
    "septimo": 7,
    "septima": 7,
    "octavo": 8,
    "octava": 8,
    "noveno": 9,
    "novena": 9,
    "décimo": 10,
    "décima": 10,
    "decimo": 10,
    "decima": 10,
    "vigésimo": 20,
    "vigésima": 20,
    "vigesimo": 20,
    "vigesima": 20,
    "trigésimo": 30,
    "trigésima": 30,
    "trigesimo": 30,
    "trigesima": 30,
    "cuadragésimo": 40,
    "cuadragésima": 40,
    "cuadragesimo": 40,
    "cuadragesima": 40,
    "quincuagésimo": 50,
    "quincuagésima": 50,
    "quincuagesimo": 50,
    "quincuagesima": 50,
    "sexagésimo": 60,
    "sexagésima": 60,
    "sexagesimo": 60,
    "sexagesima": 60,
    "septuagésimo": 70,
    "septuagésima": 70,
    "septuagesimo": 70,
    "septuagesima": 70,
    "octogésimo": 80,
    "octogésima": 80,
    "octogesimo": 80,
    "octogesima": 80,
    "nonagésimo": 90,
    "nonagésima": 90,
    "nonagesimo": 90,
    "nonagesima": 90,
    "centésimo": 100,
    "centésima": 100,
    "centesimo": 100,
    "centesima": 100,
}

# Rutas canonicas del corpus consolidado (FR-013, data-model.md:82-83).
RUTA_REGISTRO_POR_DEFECTO = "data/corpus/actos_modificatorios/.corpus_consolidado.json"
DIRECTORIO_ACTOS_POR_DEFECTO = "data/corpus/actos_modificatorios"


class ErrorIngesta(Exception):
    """Error tipificado de la ingesta de actos (data-model.md:178-191).

    Codigos propios del CLI: FORMATO_NO_SOPORTADO, SIN_TEXTO_EXTRAIBLE,
    SIN_ARTICULOS_PARSEABLES, FECHA_ANTERIOR_AL_555, FUENTE_NO_DISPONIBLE,
    DUPLICADO. La taxonomia de 10 codigos de `app/errores.py` NO se toca (FR-011).
    """

    def __init__(self, codigo: str, mensaje: str) -> None:
        super().__init__(f"Error de ingesta [{codigo}]: {mensaje}")
        self.codigo = codigo
        self.mensaje = mensaje


# --- Deteccion de formato (T004, FR-001, research D5) ---


def detectar_formato(ruta: str | Path) -> FormatoDocumento:
    """Detecta el formato del archivo fuente por magic bytes + extension.

    El contenido manda (los magic bytes son la senal autoritativa: `%PDF-` ->
    pdf, `PK\\x03\\x04` -> docx, HTML -> sisjur_html); la extension es el
    fallback. Sin formato reconocido -> ErrorIngesta FORMATO_NO_SOPORTADO
    (FR-009): el corpus existente NO se modifica (SC-006).
    """
    archivo = Path(ruta)
    if not archivo.is_file():
        raise ErrorIngesta(
            "FORMATO_NO_SOPORTADO",
            f"el archivo '{archivo.name}' no existe o no es un archivo. "
            "El corpus existente NO se modificó.",
        )
    contenido = archivo.read_bytes()[:1024]
    if contenido.startswith(b"%PDF-"):
        return "pdf"
    if contenido.startswith(b"PK\x03\x04"):
        return "docx"
    if re.search(rb"<!DOCTYPE\s+html|<html[\s>]", contenido, re.IGNORECASE):
        return "sisjur_html"
    formato = FORMATOS_POR_EXTENSION.get(archivo.suffix.lower())
    if formato is None:
        raise ErrorIngesta(
            "FORMATO_NO_SOPORTADO",
            f"el archivo '{archivo.name}' no es un formato soportado "
            "(HTML sisjur, PDF, DOCX, Markdown, TXT). El corpus existente NO se modificó.",
        )
    return formato


# --- Extraccion generica (T005, D5) ---


def _decodificar_texto(contenido: bytes) -> str:
    """Decodifica el texto con UTF-8 y fallback latin-1 (H1).

    La fuente sisjur sirve ISO-8859-1 (latin-1): sin el fallback los acentos
    del titulo/articulado romperian la decodificacion.
    """
    try:
        return contenido.decode("utf-8")
    except UnicodeDecodeError:
        return contenido.decode("latin-1")


def _texto_markdown(contenido: bytes) -> str:
    """Normaliza Markdown a texto plano: quita los marcadores de encabezado.

    "## Artículo 1." debe abrir un articulo igual que "Artículo 1.": sin la
    normalizacion, el patron generico (anclado al inicio de linea) no matchea.
    """
    texto = _decodificar_texto(contenido)
    return re.sub(r"^#{1,6}\s+", "", texto, flags=re.MULTILINE)


def _orden_lectura_reconstruido(texto: str) -> bool:
    """True si los numeros de articulo aparecen en orden no decreciente.

    pypdf colapsa el espaciado y puede desordenar columnas; pdfplumber
    reconstruye mejor el layout (D5/H6). Si el orden de aparicion de los
    numeros de articulo no es no-decreciente, el orden de lectura no se
    reconstruyo y se intenta la alternativa. IGNORECASE: el articulado real
    usa "ARTÍCULO N." en mayúsculas.
    """
    numeros = [int(n) for n in re.findall(r"Art[ií]culo\s+(\d+)", texto, re.IGNORECASE)]
    return numeros == sorted(numeros)


def _extraer_texto_pdf(contenido: bytes) -> str:
    """Extrae el texto de un PDF: pypdf primario, pdfplumber alternativa (D5).

    pypdf primero; si el orden de lectura no se reconstruye (numeros de
    articulo desordenados) se reintenta con pdfplumber. PDF sin texto extraible
    (escaneado) -> ErrorIngesta SIN_TEXTO_EXTRAIBLE sugiriendo HTML sisjur
    (edge case de la spec).
    """
    from pypdf import PdfReader

    try:
        lector = PdfReader(io.BytesIO(contenido))
        texto = "\n".join((pagina.extract_text() or "") for pagina in lector.pages)
    except Exception as e:
        raise ErrorIngesta(
            "SIN_TEXTO_EXTRAIBLE",
            "no se pudo extraer texto del PDF. Si el PDF esta escaneado, usa el "
            "formato HTML sisjur (recomendado). El corpus existente NO se modificó.",
        ) from e
    if not texto.strip():
        raise ErrorIngesta(
            "SIN_TEXTO_EXTRAIBLE",
            "el PDF no contiene texto extraible (PDF escaneado). Usa el formato "
            "HTML sisjur (recomendado). El corpus existente NO se modificó.",
        )
    if not _orden_lectura_reconstruido(texto):
        import pdfplumber

        try:
            with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                texto = "\n".join((pagina.extract_text() or "") for pagina in pdf.pages)
        except Exception as e:
            raise ErrorIngesta(
                "SIN_TEXTO_EXTRAIBLE",
                "no se pudo reconstruir el orden de lectura del PDF. Si el PDF "
                "esta escaneado, usa el formato HTML sisjur (recomendado). El "
                "corpus existente NO se modificó.",
            ) from e
        if not texto.strip():
            raise ErrorIngesta(
                "SIN_TEXTO_EXTRAIBLE",
                "el PDF no contiene texto extraible (PDF escaneado). Usa el formato "
                "HTML sisjur (recomendado). El corpus existente NO se modificó.",
            )
    return texto


def _extraer_texto_docx(contenido: bytes) -> str:
    """Extrae el texto de un DOCX: parrafos + tablas en orden de lectura (D5).

    python-docx lee los parrafos del cuerpo y las tablas intercaladas; sin
    texto -> ErrorIngesta SIN_TEXTO_EXTRAIBLE recomendando HTML sisjur.
    """
    import docx

    try:
        documento = docx.Document(io.BytesIO(contenido))
    except Exception as e:
        raise ErrorIngesta(
            "SIN_TEXTO_EXTRAIBLE",
            "no se pudo leer el DOCX. Usa el formato HTML sisjur (recomendado). "
            "El corpus existente NO se modificó.",
        ) from e
    bloques: list[str] = []
    for parrafo in documento.paragraphs:
        if parrafo.text.strip():
            bloques.append(parrafo.text)
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    bloques.append(celda.text)
    texto = "\n".join(bloques)
    if not texto.strip():
        raise ErrorIngesta(
            "SIN_TEXTO_EXTRAIBLE",
            "el DOCX no contiene texto extraible. Usa el formato HTML sisjur "
            "(recomendado). El corpus existente NO se modificó.",
        )
    return texto


def _normalizar_ordinal_textual(palabra: str) -> int | None:
    """Convierte un ordinal textual a numero (Primero -> 1, Unico -> 1).

    Devuelve None si la palabra no es un ordinal reconocido (el articulo se
    salta, no se crea con numero invalido).
    """
    clave = palabra.lower().strip(".")
    return ORDINALES_A_NUMERO.get(clave)


def _numero_articulo(partido: re.Match[str]) -> int | None:
    """Numero del articulo desde el patron generico (digito u ordinal textual)."""
    if partido.group("numero") is not None:
        return int(partido.group("numero"))
    return _normalizar_ordinal_textual(partido.group("ordinal"))


def _parsear_articulos_texto(texto: str) -> list[ArticuloNormativo]:
    """Parsea articulos de texto plano (PDF/DOCX/MD/TXT) con el patron generico.

    Cada "Artículo N." abre un articulo: el titulo es el resto de la linea del
    encabezado y el texto llega hasta el siguiente articulo (patron del parser
    demo de F2). Sin articulos parseables -> ErrorIngesta SIN_ARTICULOS_PARSEABLES
    (sin ingesta parcial, FR-009).
    """
    posiciones = list(ARTICULO_GENERICO_PATRON.finditer(texto))
    articulos: list[ArticuloNormativo] = []
    for indice, partido in enumerate(posiciones):
        numero = _numero_articulo(partido)
        if numero is None:
            continue
        fin_linea = texto.find("\n", partido.end())
        fin_encabezado = fin_linea if fin_linea != -1 else len(texto)
        # El titulo conserva su punto final (semantica del parser sisjur de F2:
        # "Objeto y ámbito de aplicación."), el patron ya consumio el ordinal.
        titulo = texto[partido.end():fin_encabezado].strip()
        inicio_cuerpo = fin_encabezado + 1 if fin_linea != -1 else len(texto)
        fin_cuerpo = (
            posiciones[indice + 1].start() if indice + 1 < len(posiciones) else len(texto)
        )
        cuerpo = texto[inicio_cuerpo:fin_cuerpo].strip()
        articulos.append(
            ArticuloNormativo(
                numero=numero,
                titulo=titulo,
                texto=cuerpo,
                libro="I",
                parte=None,
                upls_mencionadas=_extraer_upls(cuerpo),
                articulos_derogados=_extraer_derogados(cuerpo),
            )
        )
    if not articulos:
        raise ErrorIngesta(
            "SIN_ARTICULOS_PARSEABLES",
            "no se detectó ningún artículo parseable en el documento. Verifica que "
            "el documento tiene articulado o usa el formato HTML sisjur (recomendado "
            "para PDF escaneado). El corpus existente NO se modificó.",
        )
    return articulos


def extraer_articulos(contenido: bytes, formato: FormatoDocumento) -> list[ArticuloNormativo]:
    """Extrae los articulos tipados de un acto segun su formato (FR-001, D4/D5).

    sisjur_html -> `parsear_articulos` (parser de anclas de F2 con adaptacion
    D4); pdf/docx/markdown/txt -> extraccion generica (D5). Sin articulos
    parseables -> ErrorIngesta SIN_ARTICULOS_PARSEABLES (sin ingesta parcial).
    """
    if formato == "sisjur_html":
        try:
            return parsear_articulos(_decodificar_texto(contenido))
        except ValueError as e:
            raise ErrorIngesta(
                "SIN_ARTICULOS_PARSEABLES",
                f"no se detectó ningún artículo parseable en el HTML sisjur. {e} "
                "El corpus existente NO se modificó.",
            ) from e
    if formato == "pdf":
        texto = _extraer_texto_pdf(contenido)
    elif formato == "docx":
        texto = _extraer_texto_docx(contenido)
    elif formato == "markdown":
        texto = _texto_markdown(contenido)
    else:  # txt
        texto = _decodificar_texto(contenido)
    return _parsear_articulos_texto(texto)


def extraer_documento_sisjur(
    contenido: bytes,
) -> tuple[list[ArticuloNormativo], str | None, str | None]:
    """Articulos + metadatos de estado de un HTML sisjur (D4/H7).

    Devuelve (articulos, estado_documento, derogado_compilado_por): el banner
    de derogacion/compilacion de la plantilla sisjur se captura como metadato
    del documento sin romper el parseo (el banner vive fuera del articulado).
    """
    html = _decodificar_texto(contenido)
    articulos = extraer_articulos(contenido, "sisjur_html")
    estado_documento, derogado_compilado_por = _extraer_banner_derogacion(html)
    return articulos, estado_documento, derogado_compilado_por


# --- Validacion FR-014 (T008) ---


def extraer_articulos_referenciados(html: str) -> list[int]:
    """Articulos del 555 referenciados por enlaces sisjur internos (H2, FR-014).

    Deduplicacion preservando el orden de aparicion: es la referencia
    verificable por maquina para `relacion_con_555`.
    """
    return list(dict.fromkeys(int(n) for n in ENLACE_ARTICULO_555_PATRON.findall(html)))


def validar_relacion_con_555(
    fecha_expedicion: str, articulos_referenciados: list[int]
) -> str:
    """Valida FR-014 y devuelve el literal canonico de `relacion_con_555`.

    - `fecha_expedicion < FECHA_VIGENCIA_555` (2021-12-30): rechazo tipificado
      FECHA_ANTERIOR_AL_555 (fallo atomico: corpus intacto).
    - Sin referencias verificables: "sin_referencia" (se integra con warning; la
      relacion no siempre es verificable por maquina, no se rechaza).
    - Con referencias: "referencia_articulos".
    """
    if fecha_expedicion < FECHA_VIGENCIA_555:
        raise ErrorIngesta(
            "FECHA_ANTERIOR_AL_555",
            f"el acto (fecha_expedicion {fecha_expedicion}) no puede reglamentar ni "
            f"modificar el Decreto 555 de 2021 (vigencia {FECHA_VIGENCIA_555}). "
            "El corpus existente NO se modificó.",
        )
    if articulos_referenciados:
        return "referencia_articulos"
    return "sin_referencia"


# --- Registro y escritura versionada (T009, FR-013, D3) ---


def hash_archivo(contenido: bytes) -> str:
    """SHA-256 de los bytes del archivo fuente (FR-007, SC-003).

    Es la base de la deduplicacion por documento: el mismo archivo re-ingestado
    produce el mismo hash y la ingesta es no-op.
    """
    return hashlib.sha256(contenido).hexdigest()


def _escribir_atomicamente(ruta: Path, contenido: str) -> None:
    """Escribe `contenido` en `ruta` de forma atomica (FR-009, SC-006).

    Escribe a un temporal del mismo directorio y lo renombra con os.replace:
    un fallo a mitad de escritura nunca deja un archivo parcial en la ruta final.
    `mkstemp` crea el temporal con modo 0600; se fija 0644 ANTES del replace
    para que el corpus versionado sea legible por cualquier usuario del repo
    (SC-001: los JSONL/.sha256/registro salían 0600).
    """
    ruta.parent.mkdir(parents=True, exist_ok=True)
    descriptor, ruta_temporal = tempfile.mkstemp(
        dir=ruta.parent, prefix=f".{ruta.name}.", suffix=".tmp"
    )
    try:
        with os.fdopen(descriptor, "w", encoding="utf-8") as archivo:
            archivo.write(contenido)
        os.chmod(ruta_temporal, 0o644)
        os.replace(ruta_temporal, ruta)
    except Exception:
        try:
            os.unlink(ruta_temporal)
        except OSError:
            pass
        raise


def leer_registro_corpus(ruta_registro: str | Path) -> dict[str, Any]:
    """Lee el registro `.corpus_consolidado.json`; estructura vacia si no existe.

    Registro corrupto -> RuntimeError (fallo de infraestructura, FR-009): no se
    continua con un registro medio-escrito que romperia la deduplicacion.
    """
    archivo = Path(ruta_registro)
    if not archivo.exists():
        return {"documento_base": "Decreto_555_2021", "documentos": []}
    try:
        with archivo.open("r", encoding="utf-8") as f:
            registro = json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        raise RuntimeError(
            f"No se pudo leer el registro del corpus consolidado ({ruta_registro}): {e}. "
            "El corpus existente NO se modificó."
        ) from e
    if not isinstance(registro, dict):
        raise RuntimeError(
            f"El registro del corpus consolidado ({ruta_registro}) no es un objeto JSON. "
            "El corpus existente NO se modificó."
        )
    registro.setdefault("documento_base", "Decreto_555_2021")
    registro.setdefault("documentos", [])
    return registro


def marcar_documento_indexado(
    documento_id: str, ruta_registro: str | Path = RUTA_REGISTRO_POR_DEFECTO
) -> None:
    """Marca `indexado: true` de un documento en el registro tras indexar (SC-001).

    La ingesta escribe el registro con el estado de indexación conocido en ese
    momento (el `--indexar` de la llamada); si la indexación se ejecuta después
    (re-ingesta duplicada, indexado manual posterior), la entrada debe reflejar
    el estado real. No-op si el documento no está en el registro o ya está
    indexado.
    """
    registro = leer_registro_corpus(ruta_registro)
    cambiado = False
    for entrada in registro.get("documentos", []):
        if entrada.get("documento_id") == documento_id and entrada.get("indexado") is not True:
            entrada["indexado"] = True
            cambiado = True
            break
    if cambiado:
        _escribir_atomicamente(
            Path(ruta_registro),
            json.dumps(registro, ensure_ascii=False, indent=2) + "\n",
        )


def nombre_legible(documento: DocumentoNormativo) -> str:
    """Nombre legible de la norma (FR-005): 'Decreto 122 de 2023'."""
    return f"{documento.tipo_norma.capitalize()} {documento.numero} de {documento.año}"


def _articulo_a_linea(articulo: ArticuloNormativo, documento: DocumentoNormativo) -> dict[str, Any]:
    """Linea JSONL de un articulo del acto: campos F2 + metadatos de norma.

    El JSONL versionado (fuente de verdad, FR-013) conserva los campos del
    `ArticuloNormativo` de F2 y anade los campos de norma del contrato
    (contracts/ingesta-actos-modificatorios.md:86-101).
    """
    linea = articulo.model_dump()
    linea.update(
        {
            "norma_id": documento.documento_id,
            "tipo_norma": documento.tipo_norma,
            "numero_norma": documento.numero,
            "año": documento.año,
            "fecha_vigencia": documento.fecha_vigencia,
            "titulo_norma": nombre_legible(documento),
            "relacion_con_555": documento.relacion_con_555,
            "articulos_referenciados": documento.articulos_referenciados,
            "estado_documento": documento.estado_documento,
            "derogado_compilado_por": documento.derogado_compilado_por,
        }
    )
    return linea


def _entrada_registro(
    documento: DocumentoNormativo, articulos: list[ArticuloNormativo], indexado: bool
) -> dict[str, Any]:
    """Entrada del registro `.corpus_consolidado.json` por documento (FR-002/FR-007)."""
    return {
        "documento_id": documento.documento_id,
        "hash_sha256": documento.hash_sha256,
        "tipo_norma": documento.tipo_norma,
        "numero": documento.numero,
        "año": documento.año,
        "fecha_expedicion": documento.fecha_expedicion,
        "fecha_vigencia": documento.fecha_vigencia,
        "url_origen": documento.url_origen,
        "formato": documento.formato,
        "relacion_con_555": documento.relacion_con_555,
        "articulos": len(articulos),
        "indexado": indexado,
    }


def _salida_contrato(
    documento: DocumentoNormativo, articulos: list[ArticuloNormativo], duplicado: bool, indexado: bool
) -> dict[str, Any]:
    """Shape de salida en exito del CLI `acto` (contracts:38-55)."""
    return {
        "documento_id": documento.documento_id,
        "tipo_norma": documento.tipo_norma,
        "numero": documento.numero,
        "año": documento.año,
        "fecha_expedicion": documento.fecha_expedicion,
        "fecha_vigencia": documento.fecha_vigencia,
        "url_origen": documento.url_origen,
        "hash_sha256": documento.hash_sha256,
        "articulos": len(articulos),
        "relacion_con_555": documento.relacion_con_555,
        "articulos_referenciados": documento.articulos_referenciados,
        "estado_documento": documento.estado_documento,
        "derogado_compilado_por": documento.derogado_compilado_por,
        "duplicado": duplicado,
        "indexado": indexado,
    }


def escribir_documento_acto(
    contenido: bytes,
    documento: DocumentoNormativo,
    articulos: list[ArticuloNormativo],
    ruta_registro: str | Path = RUTA_REGISTRO_POR_DEFECTO,
    directorio_salida: str | Path = DIRECTORIO_ACTOS_POR_DEFECTO,
    indexado: bool = False,
) -> dict[str, Any]:
    """Registra y escribe el JSONL + `.sha256` de un acto (FR-013, D3).

    Deduplicacion por hash del archivo (FR-007, SC-003): si `hash_sha256` ya
    esta en el registro, la ingesta es no-op con `"duplicado": True` (sin
    reescribir el JSONL ni re-indexar). Escritura y registro con fallo atomico
    por documento (FR-009, SC-006): cualquier error revierte los archivos
    nuevos y deja el corpus existente intacto.

    Devuelve el shape de salida del contrato (incluye `duplicado` e `indexado`).
    """
    registro = leer_registro_corpus(ruta_registro)
    if any(
        entrada.get("hash_sha256") == documento.hash_sha256
        for entrada in registro.get("documentos", [])
    ):
        return _salida_contrato(documento, articulos, duplicado=True, indexado=False)

    directorio = Path(directorio_salida)
    ruta_jsonl = directorio / f"{documento.documento_id}.jsonl"
    ruta_sha = directorio / f"{documento.documento_id}.jsonl.sha256"

    lineas = [
        json.dumps(_articulo_a_linea(articulo, documento), ensure_ascii=False)
        for articulo in articulos
    ]
    try:
        _escribir_atomicamente(ruta_jsonl, "\n".join(lineas) + "\n")
        _escribir_atomicamente(ruta_sha, hash_documento(articulos))
        registro["documentos"].append(_entrada_registro(documento, articulos, indexado))
        _escribir_atomicamente(
            Path(ruta_registro),
            json.dumps(registro, ensure_ascii=False, indent=2) + "\n",
        )
    except Exception:
        # Fallo atomico por documento (FR-009, SC-006): si la escritura del
        # registro falla tras escribir el JSONL, se revierten los archivos
        # nuevos para dejar el corpus existente intacto.
        for ruta in (ruta_jsonl, ruta_sha):
            try:
                ruta.unlink()
            except OSError:
                pass
        raise

    return _salida_contrato(documento, articulos, duplicado=False, indexado=indexado)
