"""Fixtures compartidas: providers con respuestas simuladas via httpx.MockTransport.

Ninguna prueba hace llamadas de red reales (tasks.md: T013-T036).
"""

from __future__ import annotations

import io

import httpx

from app.main import ServidorLotes
from app.providers.arcgis import ArcGISProvider
from app.providers.mapas_bogota import MapasBogotaProvider
from app.providers.normativa import NormativaProvider
from app.providers.upl import UPLProvider


# --- Respuestas simuladas de las fuentes ---

CHIP_VALIDO = "AAA0072LRYN"
CHIP_INEXISTENTE = "ZZZ9999ZZZ9"
CODIGO_CATASTRAL = "006202003016"
MANZANA = "006202003"

RESPUESTA_CHIP_AAA = {
    "resultados": [
        {
            "OBJECTID": "68410691",
            "CODIGO_POSTAL": "111311",
            "VALUE": "AAA0072LRYN",
            "NOMBRE": "CRA 12 # 10-20",
            "BARRIO": "LAS NIEVES",
            "GEOMETRY": {
                "rings": [
                    [
                        [-74.083, 4.603],
                        [-74.082, 4.603],
                        [-74.082, 4.604],
                        [-74.083, 4.604],
                        [-74.083, 4.603],
                    ]
                ]
            },
        }
    ],
    "status": True,
}

# CHIP desconocido: la API viva responde HTTP 200 con status:false y el mensaje
# "El servicio no esta disponible" (NO es un 5xx; se mapea a "no encontrado").
RESPUESTA_CHIP_VACIA = {"mensaje": "El servicio no esta disponible", "status": False}


def geocodificar_unica():
    return {
        "resultados": [
            {"NOMBRE": "Calle 26 # 69-76", "LATITUD": 4.665, "LONGITUD": -74.102}
        ]
    }


def geocodificar_varias():
    return {
        "resultados": [
            {"NOMBRE": "Calle 26 # 69-76", "LATITUD": 4.665, "LONGITUD": -74.102},
            {"NOMBRE": "Calle 26 # 69-76 A", "LATITUD": 4.668, "LONGITUD": -74.105},
        ]
    }


def geocodificar_vacia():
    return {"resultados": []}


def feature_lote(codigo_catastral=CODIGO_CATASTRAL, manzana=MANZANA, chip=CHIP_VALIDO):
    return {
        "type": "Feature",
        "properties": {"LOTCODIGO": codigo_catastral, "MANZCODIGO": manzana, "CHIP": chip},
        "geometry": {
            "type": "Polygon",
            "coordinates": [
                [
                    [-74.084, 4.603],
                    [-74.082, 4.603],
                    [-74.082, 4.604],
                    [-74.084, 4.604],
                    [-74.084, 4.603],
                ]
            ],
        },
    }


def feature_valor(valor_m2=3200000, anio=2025):
    return {
        "type": "Feature",
        "properties": {"VALOR_M2": valor_m2, "ANIO": anio},
        "geometry": None,
    }


def feature_reserva(descripcion="Reserva vial Avenida 68", anio=None):
    propiedades = {"DESCRIPCION": descripcion}
    if anio is not None:
        propiedades["ANIO"] = anio
    return {
        "type": "Feature",
        "properties": propiedades,
        "geometry": None,
    }


def feature_obra(nombre):
    return {
        "type": "Feature",
        "properties": {"NOMBRE": nombre},
        "geometry": None,
    }


def geojson(features):
    return {"type": "FeatureCollection", "features": features}


# --- Constructores de providers ---


def provider_mapas_estandar(api_key="clave-de-prueba"):
    """Provider de Mapas Bogota: CHIP conocido, CHIP inexistente y geocodificacion unica.

    La API viva expone /buscar (cmd=direccion_chip) y /api (cmd=geocodificar)
    en https://catalogopmb.catastrobogota.gov.co/PMBWeb/web; el mock valida la
    ruta y el cmd de cada consulta.
    """

    def handler(request: httpx.Request) -> httpx.Response:
        cmd = request.url.params.get("cmd")
        if cmd == "direccion_chip" and request.url.path.endswith("/buscar"):
            query = request.url.params.get("query")
            if query == CHIP_VALIDO:
                return httpx.Response(200, json=RESPUESTA_CHIP_AAA)
            return httpx.Response(200, json=RESPUESTA_CHIP_VACIA)
        if cmd == "geocodificar" and request.url.path.endswith("/api"):
            return httpx.Response(200, json=geocodificar_unica())
        return httpx.Response(500, json={"error": "cmd no simulado"})

    return MapasBogotaProvider(transport=httpx.MockTransport(handler), api_key=api_key)


def provider_arcgis_estandar(
    lotes=None,
    valor=None,
    reserva=None,
    obras=None,
):
    """Provider ArcGIS: capa Lote por punto y las 3 tematicas activas con respuestas simuladas.

    Cada parametro acepta una lista de features o la tupla (payload, status) para
    simular errores HTTP de la fuente.
    """

    def respuesta_de(contenido):
        if isinstance(contenido, tuple) and len(contenido) == 2 and isinstance(contenido[1], int):
            payload, status = contenido
            return httpx.Response(status, json=payload)
        return httpx.Response(200, json=geojson(contenido))

    def handler(request: httpx.Request) -> httpx.Response:
        url = str(request.url)
        if "Mapa_Referencia/Mapa_Referencia/MapServer/38/query" in url:
            return respuesta_de(lotes if lotes is not None else [feature_lote()])
        if "valorreferencia" in url:
            return respuesta_de(valor if valor is not None else [feature_valor()])
        if "reservavial" in url:
            return respuesta_de(reserva if reserva is not None else [feature_reserva()])
        if "obraspublicas" in url:
            return respuesta_de(
                obras
                if obras is not None
                else [feature_obra("Parque Metropolitano"), feature_obra("Avenida Ciudad de Cali")]
            )
        if "catastro/lote/MapServer/3/query" in url:
            return respuesta_de(PAYLOAD_PREDIO)
        # Bloques F5: retornar features vacias para URLs nuevas (degradacion no_encontrado)
        if "emergencias/gestionriesgos" in url:
            return httpx.Response(200, json=geojson([]))
        if "estratificacion" in url:
            return httpx.Response(200, json=geojson([]))
        if "usopredominante" in url:
            return httpx.Response(200, json=geojson([]))
        if "alturamedia" in url:
            return httpx.Response(200, json=geojson([]))
        if "medianaavaluo" in url:
            return httpx.Response(200, json=geojson([]))
        if "licenciasconstruccion" in url:
            return httpx.Response(200, json=geojson([]))
        if "plusvalia" in url:
            return httpx.Response(200, json=geojson([]))
        if "bienesinterescultural" in url:
            return httpx.Response(200, json=geojson([]))
        if "planarqueologico" in url:
            return httpx.Response(200, json=geojson([]))
        if "transportepublico" in url:
            return httpx.Response(200, json=geojson([]))
        if "metrobogota" in url:
            return httpx.Response(200, json=geojson([]))
        # Bloques F7: catastro data
        if "catastro/construccion" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/manzana" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/densidadpredialmz" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/variacionareaconstruida" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/sectorcatastral" in url:
            return httpx.Response(200, json=geojson([]))
        return httpx.Response(404, json={"error": f"sin respuesta simulada para {url}"})

    return ArcGISProvider(transport=httpx.MockTransport(handler))


def construir_servidor(mapas=None, arcgis=None):
    """ServidorLotes con providers simulados (por defecto el flujo feliz estandar)."""
    return ServidorLotes(
        mapas if mapas is not None else provider_mapas_estandar(),
        arcgis if arcgis is not None else provider_arcgis_estandar(),
        UPLProvider(transport=httpx.MockTransport(lambda r: httpx.Response(200, json={"type": "FeatureCollection", "features": []}))),
        NormativaProvider(),
    )


# --- Fixtures Feature 3 (informe de factibilidad, T007) ---
# Patron payload/status de F1/F2: respuestas simuladas de la capa Predio
# (f=pjson), de obraspublicas con buffer 500 m (f=geojson) y de la capa UPL
# (UPL24 Chapinero, vocacion Urbano). Ninguna prueba hace llamadas de red reales.

# Capa Predio (catastro/lote/MapServer/3): formato pjson {"features": [{"attributes": {...}}]}
# (research H1/H3). CHIP AAA0072LRYN -> 2 filas: PRECDESTIN=04, PRECUSO=015/096,
# PREAUSO=40453.8/3011.3, PREVACTUAL=2026, BARMANPRE=006101016001.
PREDIO_FILA_DOMINANTE = {
    "attributes": {
        "PRECDESTIN": "04",
        "PRECUSO": "015",
        "PREAUSO": 40453.8,
        "PREVACTUAL": "2026",
        "PREATERRE": 3704.8,
        "PREACONST": 43465.1,
        "PREDIRECC": "AK 30 25 90",
        "PRENBARRIO": "FLORIDA",
        "BARMANPRE": "006101016001",
    }
}

PREDIO_FILA_SECUNDARIA = {
    "attributes": {
        "PRECDESTIN": "04",
        "PRECUSO": "096",
        "PREAUSO": 3011.3,
        "PREVACTUAL": "2026",
        "PREATERRE": 3704.8,
        "PREACONST": 43465.1,
        "PREDIRECC": "AK 30 25 90",
        "PRENBARRIO": "FLORIDA",
        "BARMANPRE": "006101016001",
    }
}

PAYLOAD_PREDIO = {"features": [PREDIO_FILA_DOMINANTE, PREDIO_FILA_SECUNDARIA]}
PAYLOAD_PREDIO_VACIO = {"features": []}

# Obras publicas con buffer 500 m (FR-004, research H5): formato geojson.
PAYLOAD_OBRAS_BUFFER_500 = geojson(
    [feature_obra("Ampliación de Estaciones: Calle 146")]
)

# Capa UPL (unidadplaneamientolocal/0): UPL24 Chapinero con vocacion Urbano.
def feature_upl(codigo="UPL24", nombre="Chapinero", vocacion="Urbano"):
    return {
        "type": "Feature",
        "properties": {"CODIGO_UPL": codigo, "NOMBRE": nombre, "VOCACION": vocacion},
        "geometry": None,
    }


def provider_arcgis_f3(lotes=None, valor=None, reserva=None, obras=None, predio=None, contador=None):
    """Provider ArcGIS del flujo F3: capa Lote, tematicas, obras buffer 500 m y capa Predio.

    `predio` acepta el payload pjson de la capa Predio o la tupla (payload, status);
    `obras` por defecto es el payload con buffer 500 m (formato geojson).
    `contador` (opcional) es una lista donde el handler registra cada request
    (str(request.url)) para verificar el numero de consultas por ruta (deuda
    tecnica post-revision: contexto tematico consultado una sola vez).
    """

    def respuesta_de(contenido):
        if isinstance(contenido, tuple) and len(contenido) == 2 and isinstance(contenido[1], int):
            payload, status = contenido
            return httpx.Response(status, json=payload)
        if isinstance(contenido, list):
            return httpx.Response(200, json=geojson(contenido))
        return httpx.Response(200, json=contenido)

    def handler(request: httpx.Request) -> httpx.Response:
        if contador is not None:
            contador.append(str(request.url))
        url = str(request.url)
        if "Mapa_Referencia/Mapa_Referencia/MapServer/38/query" in url:
            return respuesta_de(lotes if lotes is not None else [feature_lote(codigo_catastral="006101016001", manzana="006101016")])
        if "valorreferencia" in url:
            return respuesta_de(valor if valor is not None else [feature_valor()])
        if "reservavial" in url:
            return respuesta_de(reserva if reserva is not None else [feature_reserva()])
        if "obraspublicas" in url:
            return respuesta_de(obras if obras is not None else PAYLOAD_OBRAS_BUFFER_500)
        if "catastro/lote/MapServer/3/query" in url:
            return respuesta_de(predio if predio is not None else PAYLOAD_PREDIO)
        # Bloques F5: retornar features vacias para URLs nuevas (degradacion no_encontrado)
        if "emergencias/gestionriesgos" in url:
            return httpx.Response(200, json=geojson([]))
        if "estratificacion" in url:
            return httpx.Response(200, json=geojson([]))
        if "usopredominante" in url:
            return httpx.Response(200, json=geojson([]))
        if "alturamedia" in url:
            return httpx.Response(200, json=geojson([]))
        if "medianaavaluo" in url:
            return httpx.Response(200, json=geojson([]))
        if "licenciasconstruccion" in url:
            return httpx.Response(200, json=geojson([]))
        if "plusvalia" in url:
            return httpx.Response(200, json=geojson([]))
        if "bienesinterescultural" in url:
            return httpx.Response(200, json=geojson([]))
        if "planarqueologico" in url:
            return httpx.Response(200, json=geojson([]))
        if "transportepublico" in url:
            return httpx.Response(200, json=geojson([]))
        if "metrobogota" in url:
            return httpx.Response(200, json=geojson([]))
        # Bloques F7: catastro data
        if "catastro/construccion" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/manzana" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/densidadpredialmz" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/variacionareaconstruida" in url:
            return httpx.Response(200, json=geojson([]))
        if "catastro/sectorcatastral" in url:
            return httpx.Response(200, json=geojson([]))
        return httpx.Response(404, json={"error": f"sin respuesta simulada para {url}"})

    return ArcGISProvider(transport=httpx.MockTransport(handler))


def provider_upl_estandar(upl_features=None):
    """Provider UPL: UPL24 Chapinero (vocacion Urbano) por defecto."""
    features = upl_features if upl_features is not None else [feature_upl()]
    return UPLProvider(
        transport=httpx.MockTransport(lambda r: httpx.Response(200, json=geojson(features)))
    )


def server_lotes_f3(mapas=None, arcgis=None, upl=None, normativa=None):
    """ServidorLotes con providers simulados del flujo F3 (informe de factibilidad)."""
    return ServidorLotes(
        mapas if mapas is not None else provider_mapas_estandar(),
        arcgis if arcgis is not None else provider_arcgis_f3(),
        upl if upl is not None else provider_upl_estandar(),
        normativa if normativa is not None else NormativaProvider(),
    )


# --- Ampliacion de fixtures F3: stub del NormativaProvider (T015 y reporte F3) ---
# Los contract tests F3 no pueden usar el NormativaProvider real (requiere
# ChromaDB + Ollama, prohibido en tests). Este stub implementa la interfaz que
# usara el orquestador de `get_feasibility_report` (consultar + aclose), registra
# las llamadas para verificar la consulta explicita/automatica (T015) e inyecta
# errores tipados (CorpusNoIngestadoError/OllamaNoDisponibleError) o respuestas
# vacias para probar la degradacion deliberada de normative_evidence (FR-009).


def respuesta_normativa_ok():
    """Respuesta RAG con 1 articulo recuperado (formato consultar_normativa de F2).

    `trazabilidad` es el source_trace del corpus (Decreto 555/2021) que el
    orquestador debe propagar al bloque normative_evidence (T017).
    """
    return {
        "respuesta": "El Artículo 361 regula los usos del suelo.",
        "sin_resultados": False,
        "resultados": [
            {
                "articulo": 361,
                "titulo": "Usos del suelo",
                "libro": "III",
                "parte": "urbano",
                "texto_cita": "El presente artículo regula los usos del suelo...",
                "similitud": 0.42,
            }
        ],
        "trazabilidad": {
            "source_name": "Decreto 555 de 2021 (POT Bogotá)",
            "layer_id": "Decreto_555_2021",
            "service_url": "https://sisjur.bogotajuridica.gov.co/sisjur/normas/Norma1.jsp?i=119582",
            "data_vigencia": "2021-12-30",
            "query_timestamp": "2026-08-12T02:15:04Z",
        },
    }


def respuesta_normativa_sin_resultados():
    """Respuesta RAG sin resultados (formato consultar_normativa de F2)."""
    return {
        "respuesta": "No se encontraron resultados relevantes en el POT 555/2021.",
        "sin_resultados": True,
        "resultados": [],
        "trazabilidad": {
            "source_name": "Decreto 555 de 2021 (POT Bogotá)",
            "layer_id": "Decreto_555_2021",
            "service_url": "https://sisjur.bogotajuridica.gov.co/sisjur/normas/Norma1.jsp?i=119582",
            "data_vigencia": "2021-12-30",
            "query_timestamp": "2026-08-12T02:15:04Z",
        },
    }


class NormativaProviderStub:
    """Stub del NormativaProvider para los contract tests F3 (sin red ni Ollama).

    Registra cada llamada a `consultar` en `llamadas` (consulta, upl, top_k) y
    devuelve `respuesta` o lanza `error` si se inyecto uno (degradacion T015).
    """

    def __init__(
        self,
        respuesta: dict | None = None,
        error: Exception | None = None,
    ) -> None:
        self.respuesta = respuesta if respuesta is not None else respuesta_normativa_sin_resultados()
        self.error = error
        self.llamadas: list[dict] = []

    async def consultar(self, consulta: str, upl: str | None = None, top_k: int = 3) -> dict:
        self.llamadas.append({"consulta": consulta, "upl": upl, "top_k": top_k})
        if self.error is not None:
            raise self.error
        return self.respuesta

    async def aclose(self) -> None:
        return None


# --- Fixtures Feature 4 (ingesta de actos modificatorios, T010) ---
# Decreto 122 de 2023 (research H2/H7): reglamenta los artículos 233, 243 y 384
# del Decreto 555 de 2021 (vivienda colectiva). Se sirve en los 5 formatos
# soportados (HTML sisjur, PDF, DOCX, Markdown, TXT) con el mismo articulado:
# 13 artículos con ordinal `Nº.`. Ninguna prueba hace llamadas de red reales.

DECRETO_122_TITULOS = [
    "Objeto y ámbito de aplicación.",
    "Definiciones.",
    "Reglas generales de la vivienda colectiva.",
    "Soluciones habitacionales con servicios.",
    "Estándares mínimos de vivienda VIS y VIP.",
    "Cesiones urbanísticas.",
    "Condiciones de edificabilidad.",
    "Procedimiento de aprobación.",
    "Obligaciones de los promotores.",
    "Régimen de transición.",
    "Seguimiento y control.",
    "Sanciones.",
    "Vigencia.",
]

# Metadatos canónicos del acto (H4/H7): el Decreto 122 se expidió tras la
# vigencia del 555 (2021-12-30), referencia los artículos 233/243/384 y el
# banner sisjur lo marca derogado/compilado por el DUDOT 670 de 2025.
DECRETO_122_METADATA = {
    "tipo_norma": "decreto",
    "numero": 122,
    "año": 2023,
    "documento_id": "Decreto_122_2023",
    "fecha_expedicion": "2023-03-30",
    "fecha_vigencia": "2023-03-31",
    "url_origen": "https://www.alcaldiabogota.gov.co/sisjur/normas/Norma1.jsp?i=139499",
}

BANNER_DEROGACION_122 = (
    "Derogado y compilado por el art. 1526, Decreto Único Distrital de "
    "Ordenamiento Territorial 670 de 2025"
)


def _parrafo_articulo_decreto_122(numero, titulo, cuerpo):
    """Párrafo sisjur de un artículo del Decreto 122 (plantilla D4/H2).

    `<span class="ancla" id="N">` marca el artículo; el número es un ordinal
    (`<b>1º.</b>`, no el punto plano del 555) y el título vive en la variante
    `<i style="font-weight: bold;">` en lugar del `<b>` del 555.
    """
    return (
        f'<p class="MsoNormal"><b>Artículo</b><span style="font-size: 12pt;" '
        f'class="ancla" id="{numero}"></span>\n'
        f"<b>{numero}º.</b>&nbsp;<i style=\"font-weight: bold;\">{titulo}</i> {cuerpo}</p>"
    )


def html_decreto_122():
    """HTML sisjur del Decreto 122: 13 anclas y el banner de derogación.

    El artículo 1 enlaza los artículos 233/243/384 del 555 vía
    `Norma1.jsp?i=119582#NNN` (H2, referencia verificable por máquina para
    `relacion_con_555`); el banner vive FUERA de los `<p class="MsoNormal">`
    del articulado (H7), como en la plantilla real.
    """
    articulos = []
    for i, titulo in enumerate(DECRETO_122_TITULOS, start=1):
        if i == 1:
            cuerpo = (
                "El presente decreto reglamenta los artículos "
                '<a href="../normas/Norma1.jsp?i=119582#233">233</a>, '
                '<a href="../normas/Norma1.jsp?i=119582#243">243</a> y '
                '<a href="../normas/Norma1.jsp?i=119582#384">384</a> del Decreto '
                "Distrital 555 de 2021, en lo relacionado con la vivienda colectiva."
            )
        else:
            cuerpo = f"Cuerpo del artículo {i} del Decreto 122."
        articulos.append(_parrafo_articulo_decreto_122(i, titulo, cuerpo))
    return (
        "<html><head><title>Decreto 122 de 2023</title></head><body>\n"
        f'<div class="banner">{BANNER_DEROGACION_122}</div>\n'
        + "\n".join(articulos)
        + "\n</body></html>"
    )


def texto_decreto_122():
    """Texto plano del Decreto 122: 'ARTÍCULO Nº. Título' + cuerpo por línea.

    Es la fuente del formato TXT y el contenido incrustado en PDF y DOCX; la
    variante Markdown solo añade los marcadores `##` que `_texto_markdown`
    elimina antes de parsear.
    """
    lineas = []
    for i, titulo in enumerate(DECRETO_122_TITULOS, start=1):
        if i == 1:
            cuerpo = (
                "El presente decreto reglamenta los artículos 233, 243 y 384 del "
                "Decreto Distrital 555 de 2021, en lo relacionado con la vivienda "
                "colectiva."
            )
        else:
            cuerpo = f"Cuerpo del artículo {i} del Decreto 122."
        lineas.append(f"ARTÍCULO {i}º. {titulo}")
        lineas.append(cuerpo)
    return "\n".join(lineas)


def md_decreto_122():
    """Markdown del Decreto 122: los encabezados `## ARTÍCULO Nº.`."""
    return "\n".join(
        f"## {linea}" if linea.startswith("ARTÍCULO") else linea
        for linea in texto_decreto_122().split("\n")
    )


def txt_decreto_122():
    """TXT del Decreto 122 (idéntico al texto plano)."""
    return texto_decreto_122()


def _pdf_con_texto(texto):
    """Construye un PDF mínimo con `texto` (contenido latin-1) y xref válida.

    Una página con una fuente Type1 Helvetica (WinAnsi): cada línea del texto
    se muestra con `Tj` y se baja con `Td` para que pypdf reconstruya los
    saltos de línea y `_parsear_articulos_texto` pueda anclar los artículos.
    """
    lineas = texto.split("\n")
    operadores = " ".join(f"({linea}) Tj 0 -14 Td" for linea in lineas)
    contenido = f"BT /F1 10 Tf 72 720 Td {operadores} ET".encode("latin-1")
    objetos = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(contenido)).encode()
        + b" >>\nstream\n"
        + contenido
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
    ]
    cuerpo = bytearray(b"%PDF-1.4\n")
    offsets = []
    for indice, objeto in enumerate(objetos, start=1):
        offsets.append(len(cuerpo))
        cuerpo.extend(f"{indice} 0 obj\n".encode())
        cuerpo.extend(objeto)
        cuerpo.extend(b"\nendobj\n")
    posicion_xref = len(cuerpo)
    cuerpo.extend(f"xref\n0 {len(objetos) + 1}\n".encode())
    cuerpo.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        cuerpo.extend(f"{offset:010d} 00000 n \n".encode())
    cuerpo.extend(
        f"trailer\n<< /Size {len(objetos) + 1} /Root 1 0 R >>\nstartxref\n{posicion_xref}\n%%EOF\n".encode()
    )
    return bytes(cuerpo)


def pdf_decreto_122():
    """PDF mínimo con el articulado del Decreto 122 (texto latin-1)."""
    return _pdf_con_texto(texto_decreto_122())


def docx_decreto_122():
    """DOCX del Decreto 122: un párrafo por encabezado y por cuerpo.

    python-docx se importa perezosamente (dependencia opcional `ingesta`);
    los tests que no tocan F4 no cargan el módulo.
    """
    import docx

    documento = docx.Document()
    for i, titulo in enumerate(DECRETO_122_TITULOS, start=1):
        if i == 1:
            cuerpo = (
                "El presente decreto reglamenta los artículos 233, 243 y 384 del "
                "Decreto Distrital 555 de 2021, en lo relacionado con la vivienda "
                "colectiva."
            )
        else:
            cuerpo = f"Cuerpo del artículo {i} del Decreto 122."
        documento.add_paragraph(f"ARTÍCULO {i}º. {titulo}")
        documento.add_paragraph(cuerpo)
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


# --- Registro del corpus consolidado para F4 (T010) ---
# El registro `.corpus_consolidado.json` (FR-013, data-model.md:82-83) declara
# el documento base y la lista de documentos; `entrada_registro_corpus` arma
# una entrada canónica del Decreto 122 para inyectar duplicados o estados
# previos sin tocar `data/` real.

REGISTRO_CORPUS_PRUEBA = {
    "documento_base": "Decreto_555_2021",
    "documentos": [],
}


def entrada_registro_corpus(**campos):
    """Entrada canónica de registro del Decreto 122 (overridable por `campos`).

    `hash_sha256` por defecto es ficticio ('abc123'): los tests que simulan un
    duplicado inyectan el hash del archivo de la fixture para que
    `escribir_documento_acto` lo detecte como ya indexado (FR-007).
    """
    entrada = {
        "documento_id": DECRETO_122_METADATA["documento_id"],
        "hash_sha256": "abc123",
        "tipo_norma": DECRETO_122_METADATA["tipo_norma"],
        "numero": DECRETO_122_METADATA["numero"],
        "año": DECRETO_122_METADATA["año"],
        "fecha_expedicion": DECRETO_122_METADATA["fecha_expedicion"],
        "fecha_vigencia": DECRETO_122_METADATA["fecha_vigencia"],
        "url_origen": DECRETO_122_METADATA["url_origen"],
        "formato": "sisjur_html",
        "relacion_con_555": "referencia_articulos",
        "articulos": 13,
        "indexado": False,
    }
    entrada.update(campos)
    return entrada
